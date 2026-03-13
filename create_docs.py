from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_full_deployment_docx():
    doc = Document()

    # --- Styles Setup ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    h1_style = doc.styles['Heading 1']
    h1_style.font.color.rgb = RGBColor(46, 116, 181) # Corporate Blue
    h1_style.font.size = Pt(16)

    h2_style = doc.styles['Heading 2']
    h2_style.font.color.rgb = RGBColor(68, 114, 196)
    h2_style.font.size = Pt(13)

    # --- Title Page ---
    doc.add_heading('Полное руководство по развертыванию проекта Gide (Inturex)', 0)
    
    p = doc.add_paragraph('Версия документа: 2.0 (Расширенная)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('Дата: 09 Января 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('_' * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- 1. Архитектура ---
    doc.add_heading('1. Архитектура решения', level=1)
    doc.add_paragraph('Проект разворачивается в изолированной среде Docker, чтобы не конфликтовать с другими сайтами на сервере.')
    
    p = doc.add_paragraph()
    p.add_run('Состав контейнеров (docker-compose):').bold = True
    
    items = [
        ('gide_frontend', 'React приложение (порт 80 внутри сети).'),
        ('gide_backend', 'Python FastAPI сервер (порт 8000 внутри сети).'),
        ('gide_nginx', 'Внутренний прокси, маршрутизирует /api на бэкенд и статику на фронтенд (порт 8080 наружу).'),
        ('gide_webhook', 'Сервис автодеплоя, слушает Github (порт 9000 наружу).')
    ]
    for title, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(title).bold = True
        p.add_run(f': {desc}')

    doc.add_paragraph('\nСхема трафика:')
    doc.add_paragraph('Internet -> Host Nginx (80/443) -> Docker Nginx (8080) -> App Containers')

    # --- 2. Подготовка сервера ---
    doc.add_heading('2. Подготовка сервера', level=1)
    doc.add_paragraph('Если сервер чистый, выполните следующие команды. Если Docker уже установлен, переходите к шагу 3.')

    doc.add_heading('2.1. Установка Docker и Compose', level=2)
    
    code = """# Обновляем пакеты
sudo apt-get update && sudo apt-get upgrade -y

# Устанавливаем Docker
curl -fsSL https://get.docker.com -o get-docker.sh
sudo sh get-docker.sh

# Проверяем установку
docker --version
docker compose version"""
    
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = code
    cell.paragraphs[0].runs[0].font.name = 'Courier New'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    
    # --- 3. Установка проекта ---
    doc.add_heading('3. Установка и настройка проекта', level=1)

    doc.add_heading('Шаг 1: Клонирование репозитория', level=2)
    doc.add_paragraph('Создаем рабочую директорию и скачиваем код:')
    
    code = """mkdir -p /var/www/gide
cd /var/www/gide
git clone https://github.com/kbeautyg/gide.git ."""
    
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = code

    doc.add_heading('Шаг 2: Конфигурация (.env)', level=2)
    doc.add_paragraph('Создайте файл .env:')
    doc.add_paragraph('nano .env', style='Quote')
    doc.add_paragraph('Вставьте следующие данные (актуальные для продакшна):')

    env_content = """# --- Backend ---
DATABASE_URL=postgresql://supabase_admin:kkgw0gaylup95bsdlkjtlh90j1rj7kuaatzy22jgyv5tgogghvo9cmf7zpmwrxkx@yamabiko.proxy.rlwy.net:36914/postgres

SECRET_KEY_BASE=a1b2c3d4e5f6g7h8i9j0k1l2m3n4o5p6q7r8s9t0u1v2w3x4y5z6a7b8c9d0e1f2g3h4i5j6k7l8m9n0o1p2q3r4s5t6u7v8w9x0y1z2
SUPABASE_KEY=eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJyb2xlIjoic2VydmljZV9yb2xlIiwiaXNzIjoic3VwYWJhc2UiLCJpYXQiOjE3NDYxMzMyMDAsImV4cCI6MTkwMzg5OTYwMH0.LRUaNwqp5qFamFjI81ibwPZn75UtMK-odWFkRMAYyt0

GUIDE_PHONE=+79932890755
SUPER_ADMIN_PHONE=+79177445182
TELEGRAM_BOT_TOKEN=8409730364:AAF1NGhtiQaKkh_5QLi9DjFhgBUnVOosvUA
WEBHOOK_SECRET=my_webhook_secret

# ВАЖНО: Замените на реальный домен
BACKEND_CORS_ORIGINS=["https://gide.ваш-домен.ru"]
APP_BACKEND_URL=https://gide.ваш-домен.ru/api

# --- Frontend ---
VITE_API_URL=https://gide.ваш-домен.ru/api/v1"""

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = env_content
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0, 100, 0)

    doc.add_heading('Шаг 3: Запуск контейнеров', level=2)
    doc.add_paragraph('Выдаем права скрипту и запускаем сборку:')
    
    code = """chmod +x deploy.sh
docker-compose -f docker-compose.prod.yml up --build -d"""
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = code

    doc.add_paragraph('После выполнения команды (может занять 2-3 минуты) проверьте статус:')
    doc.add_paragraph('docker ps', style='Quote')
    doc.add_paragraph('Вы должны увидеть 4 активных контейнера (gide_frontend, gide_backend, gide_nginx, gide_webhook).')

    doc.add_page_break()

    # --- 4. Настройка Nginx Host ---
    doc.add_heading('4. Настройка системного Nginx', level=1)
    doc.add_paragraph('Чтобы сайт был доступен из интернета, нужно настроить проксирование на хосте.')

    doc.add_paragraph('1. Создайте конфиг:')
    doc.add_paragraph('nano /etc/nginx/sites-available/gide.conf', style='Quote')

    doc.add_paragraph('2. Вставьте содержимое (замените домен):')
    nginx_conf = """server {
    listen 80;
    server_name gide.ваш-домен.ru; # <-- ВАШ ДОМЕН

    # Логи (опционально)
    access_log /var/log/nginx/gide_access.log;
    error_log /var/log/nginx/gide_error.log;

    # Основное приложение (Frontend + API)
    # Проксируем на порт 8080 (Docker)
    location / {
        proxy_pass http://127.0.0.1:8080;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;
        proxy_set_header X-Forwarded-Proto $scheme;
    }
    
    # Вебхук для автодеплоя
    # Проксируем на порт 9000 (Docker Webhook Service)
    location /webhook {
        proxy_pass http://127.0.0.1:9000/webhook;
        proxy_set_header Host $host;
        proxy_set_header X-Real-IP $remote_addr;
    }
}"""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = nginx_conf
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    doc.add_paragraph('\n3. Активируйте сайт и получите SSL:')
    code = """ln -s /etc/nginx/sites-available/gide.conf /etc/nginx/sites-enabled/
nginx -t
systemctl reload nginx
certbot --nginx -d gide.ваш-домен.ru"""
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = code

    # --- 5. Обслуживание ---
    doc.add_heading('5. Обслуживание и диагностика', level=1)
    
    doc.add_heading('Просмотр логов', level=2)
    doc.add_paragraph('Если что-то не работает, проверьте логи контейнеров:')
    doc.add_paragraph('cd /var/www/gide\ndocker-compose -f docker-compose.prod.yml logs -f --tail=50', style='Quote')

    doc.add_heading('Ручное обновление', level=2)
    doc.add_paragraph('Если автодеплой не сработал, можно обновить вручную:')
    doc.add_paragraph('./deploy.sh', style='Quote')

    doc.add_heading('Перезагрузка', level=2)
    doc.add_paragraph('docker-compose -f docker-compose.prod.yml restart', style='Quote')

    doc.add_heading('Проверка работы вебхука', level=2)
    doc.add_paragraph('Можно отправить тестовый запрос из консоли сервера, чтобы проверить, работает ли автодеплой:')
    doc.add_paragraph('curl -X POST http://localhost:9000/webhook', style='Quote')

    doc.save('ADMIN_DEPLOY_DOCS_FULL.docx')
    print('ADMIN_DEPLOY_DOCS_FULL.docx created successfully')

if __name__ == '__main__':
    create_full_deployment_docx()
