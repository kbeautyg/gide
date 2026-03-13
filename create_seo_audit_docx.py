# -*- coding: utf-8 -*-
"""
Генератор SEO аудита в формате DOCX
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Устанавливает цвет фона ячейки"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, col_widths=None):
    """Добавляет таблицу в документ"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Заголовки
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], 'E8E8E8')
    
    # Данные
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
    
    return table

def create_audit():
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # ============ ТИТУЛЬНАЯ СТРАНИЦА ============
    title = doc.add_heading('ТЕХНИЧЕСКИЙ SEO АУДИТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Сайт: Inturex — Экскурсии по Азии')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    url = doc.add_paragraph('https://inturex.pro/')
    url.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph(f'Дата аудита: {datetime.now().strftime("%d.%m.%Y")}')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tech = doc.add_paragraph('Технологии: React + Vite, FastAPI, React-Helmet-Async')
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ============ СОДЕРЖАНИЕ ============
    doc.add_heading('СОДЕРЖАНИЕ', level=1)
    
    contents = [
        '1. Meta-теги и заголовки',
        '2. Open Graph теги',
        '3. Twitter Cards',
        '4. JSON-LD структурированные данные',
        '5. Hreflang и интернационализация',
        '6. Robots.txt',
        '7. Sitemap.xml',
        '8. Изображения',
        '9. Безопасность и внешние ссылки',
        '10. Доступность (A11Y)',
        '11. PWA и Manifest',
        '12. Favicons',
        '13. Технические файлы',
        '14. Core Web Vitals факторы',
        '15. Внутренняя перелинковка',
        '16. Индексация',
        '17. Структура URL',
        '18. Количественные показатели',
        '19. Сводная таблица статусов'
    ]
    
    for item in contents:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ============ 1. META-ТЕГИ ============
    doc.add_heading('1. META-ТЕГИ И ЗАГОЛОВКИ', level=1)
    
    doc.add_heading('1.1 Главная страница (index.html)', level=2)
    
    add_table(doc, ['Элемент', 'Значение', 'Статус'], [
        ['<html lang>', 'ru', '✓ Установлен'],
        ['<meta charset>', 'UTF-8', '✓ Установлен'],
        ['<meta viewport>', 'width=device-width, initial-scale=1.0', '✓ Корректный'],
        ['<title>', 'Экскурсии по Азии 2025 — Таиланд, Вьетнам, Китай, Япония | Inturex (69 символов)', '✓ Оптимальная длина'],
        ['<meta name="title">', 'Дублирует <title>', '✓ Присутствует'],
        ['<meta name="description">', '172 символа с эмодзи', '✓ Присутствует'],
        ['<meta name="keywords">', '10 ключевых слов', '✓ Присутствует'],
        ['<meta name="author">', 'Inturex', '✓ Установлен'],
        ['<meta name="robots">', 'index, follow, max-image-preview:large, max-snippet:-1, max-video-preview:-1', '✓ Полная настройка'],
        ['<link rel="canonical">', 'https://inturex.pro/', '✓ Установлен'],
        ['<meta name="theme-color">', '#6366f1', '✓ Установлен'],
        ['<meta name="format-detection">', 'telephone=no', '✓ Установлен'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('1.2 Страницы с динамическими meta-тегами (React-Helmet)', level=2)
    
    add_table(doc, ['Страница', 'Title', 'Description', 'Keywords', 'Canonical'], [
        ['HomePage', 'Динамический', 'Динамический', 'Динамический', '✓'],
        ['ToursPage', 'Динамический (зависит от location)', 'Динамический', 'Динамический', 'Динамический'],
        ['TourDetailPage', '{tour.title} — Экскурсия в {location} | Inturex', 'Первые 160 символов описания', 'Динамический', 'Динамический'],
        ['AboutPage', 'Статический', 'Статический', 'Статический', '✓ /about'],
        ['ContactPage', 'Статический', 'Статический', 'Статический', '✓ /contact'],
        ['FAQPage', 'Статический', 'Статический', 'Статический', '✓ /faq'],
        ['JournalPage', 'Отсутствует Helmet', '—', '—', '—'],
        ['ArticlePage', 'Динамический', 'Динамический', 'Отсутствует', '✓ Динамический'],
    ])
    
    doc.add_page_break()
    
    # ============ 2. OPEN GRAPH ============
    doc.add_heading('2. OPEN GRAPH ТЕГИ', level=1)
    
    doc.add_heading('2.1 Глобальные OG-теги (index.html)', level=2)
    
    add_table(doc, ['Тег', 'Значение', 'Статус'], [
        ['og:type', 'website', '✓'],
        ['og:url', 'Абсолютный URL', '✓'],
        ['og:title', 'Дублирует title', '✓'],
        ['og:description', 'Краткое описание', '✓'],
        ['og:image', '/og-image.jpg (абсолютный)', '✓'],
        ['og:locale', 'ru_RU', '✓'],
        ['og:site_name', 'Inturex — Экскурсии по Азии', '✓'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('2.2 Страничные OG-теги', level=2)
    
    add_table(doc, ['Страница', 'og:type', 'og:title', 'og:description', 'og:image', 'og:url'], [
        ['HomePage', 'website', '✓', '✓', '✓', '✓'],
        ['ToursPage', 'website', 'Динамический', 'Динамический', 'Статический', 'Динамический'],
        ['TourDetailPage', 'product', 'Динамический', 'Динамический', 'Первое фото тура', 'Динамический'],
        ['AboutPage', 'website', '✓', '✓', '✓', '✓'],
        ['ContactPage', 'website', '✓', '✓', '✓', '✓'],
        ['FAQPage', 'website', '✓', '✓', '✓', '✓'],
        ['ArticlePage', 'article', '✓', '✓', 'Фото статьи', '✓'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('2.3 Дополнительные OG-теги (TourDetailPage)', level=2)
    
    add_table(doc, ['Тег', 'Значение'], [
        ['product:price:amount', 'Цена тура'],
        ['product:price:currency', 'RUB'],
    ])
    
    doc.add_page_break()
    
    # ============ 3. TWITTER CARDS ============
    doc.add_heading('3. TWITTER CARDS', level=1)
    
    add_table(doc, ['Элемент', 'Значение', 'Статус'], [
        ['twitter:card', 'summary_large_image', '✓'],
        ['twitter:title', 'Дублирует OG', '✓'],
        ['twitter:description', 'Дублирует OG', '✓'],
        ['twitter:image', '/og-image.jpg', '✓'],
        ['twitter:url', 'Дублирует OG', '✓'],
    ])
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Примечание: ').bold = True
    p.add_run('Используется атрибут "property" вместо "name" для Twitter тегов — это работает, но не соответствует официальной спецификации.')
    
    doc.add_page_break()
    
    # ============ 4. JSON-LD ============
    doc.add_heading('4. JSON-LD СТРУКТУРИРОВАННЫЕ ДАННЫЕ', level=1)
    
    doc.add_heading('4.1 Глобальные схемы (index.html)', level=2)
    
    add_table(doc, ['Схема', 'Содержание', 'Статус'], [
        ['TravelAgency', 'name, description, url, logo, address, areaServed (11 стран), priceRange, openingHoursSpecification', '✓ Полная'],
        ['WebSite', 'name, url, potentialAction (SearchAction)', '✓ С поиском'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('4.2 Страничные схемы', level=2)
    
    add_table(doc, ['Страница', 'Схемы', 'Статус'], [
        ['HomePage', 'TravelAgency, ItemList (11 направлений)', '✓'],
        ['ToursPage', 'BreadcrumbList (при location), ItemList (топ-10 туров)', '✓'],
        ['TourDetailPage', 'TouristTrip, BreadcrumbList, Product (с отзывами)', '✓ Полная'],
        ['AboutPage', 'AboutPage, Organization', '✓'],
        ['ContactPage', 'ContactPage, Organization с ContactPoint', '✓'],
        ['FAQPage', 'FAQPage (18 вопросов)', '✓'],
        ['JournalPage', 'Отсутствует', '✗'],
        ['ArticlePage', 'Отсутствует', '✗'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('4.3 Детали схемы TouristTrip (TourDetailPage)', level=2)
    
    add_table(doc, ['Поле', 'Наличие'], [
        ['name', '✓'],
        ['description', '✓'],
        ['image', '✓ (массив)'],
        ['url', '✓'],
        ['touristType', '✓'],
        ['provider', '✓ TravelAgency'],
        ['offers', '✓ с ценой в RUB'],
        ['itinerary', '✓ (при наличии location)'],
        ['aggregateRating', '✓ (при наличии rating)'],
        ['duration', '✓ в формате ISO 8601'],
    ])
    
    doc.add_page_break()
    
    # ============ 5. HREFLANG ============
    doc.add_heading('5. HREFLANG И ИНТЕРНАЦИОНАЛИЗАЦИЯ', level=1)
    
    add_table(doc, ['Тег', 'Значение', 'Статус'], [
        ['<link hreflang="ru">', 'Основной URL', '✓'],
        ['<link hreflang="x-default">', 'Тот же URL', '✓'],
        ['Другие языки', 'Отсутствуют', '—'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('5.1 Geo-теги', level=2)
    
    add_table(doc, ['Тег', 'Значение'], [
        ['geo.region', 'TH'],
        ['geo.placename', 'Bangkok, Thailand'],
    ])
    
    doc.add_page_break()
    
    # ============ 6. ROBOTS.TXT ============
    doc.add_heading('6. ROBOTS.TXT', level=1)
    
    p = doc.add_paragraph()
    p.add_run('Расположение: ').bold = True
    p.add_run('/public/robots.txt + динамический /api/v1/seo/robots.txt')
    
    doc.add_paragraph()
    
    add_table(doc, ['Директива', 'Значение'], [
        ['User-agent', '*'],
        ['Allow', '/'],
        ['Disallow', '/dashboard/, /admin/, /login, /register, /api/'],
        ['Sitemap', 'Динамический + статический'],
        ['Crawl-delay', '1'],
        ['Host (Yandex)', '✓ Установлен'],
    ])
    
    doc.add_page_break()
    
    # ============ 7. SITEMAP.XML ============
    doc.add_heading('7. SITEMAP.XML', level=1)
    
    doc.add_heading('7.1 Статический sitemap (/public/sitemap.xml)', level=2)
    
    add_table(doc, ['Категория', 'Количество URL', 'Priority'], [
        ['Главная', '1', '1.0'],
        ['Каталог туров', '1', '0.9'],
        ['Страны', '11', '0.8'],
        ['Города', '32', '0.7'],
        ['Информационные', '8', '0.3-0.7'],
        ['ИТОГО', '53', '—'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('7.2 Динамический sitemap (/api/v1/seo/sitemap.xml)', level=2)
    
    add_table(doc, ['Категория', 'Содержание'], [
        ['Страны', '11 (жёстко закодированы)'],
        ['Города', '45 популярных'],
        ['Туры', 'Все активные туры из БД'],
        ['Информационные', '8 страниц'],
        ['lastmod', '✓ Для туров — дата обновления'],
        ['changefreq', '✓ Установлен для всех'],
        ['priority', '✓ Иерархический'],
    ])
    
    doc.add_page_break()
    
    # ============ 8. ИЗОБРАЖЕНИЯ ============
    doc.add_heading('8. ИЗОБРАЖЕНИЯ', level=1)
    
    doc.add_heading('8.1 Атрибуты изображений', level=2)
    
    add_table(doc, ['Компонент', 'alt', 'loading="lazy"', 'Статус'], [
        ['TourCard', '✓ {tour.title}', '✓', 'Корректно'],
        ['HomePage (countries)', '✓ {country.name}', '✓ (для moreCountries)', 'Частично'],
        ['HomePage (cities)', '✓ {city.name}', '✓', 'Корректно'],
        ['AboutPage', '✓ {member.name}', '✓', 'Корректно'],
        ['JournalPage', '✓ {article.title}', '✓', 'Корректно'],
        ['ArticlePage', '✓ {article.title}', '✓ eager для hero', 'Корректно'],
        ['CountryCard', '✓', '✓', 'Корректно'],
        ['CityCard', '✓', '✓', 'Корректно'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('8.2 Оптимизация изображений', level=2)
    
    add_table(doc, ['Аспект', 'Статус'], [
        ['WebP формат', '✗ Не используется'],
        ['srcset / sizes', '✗ Не используется'],
        ['Preload критических', '✗ Не настроено'],
        ['CDN для изображений', '✓ Unsplash, Tripster'],
    ])
    
    doc.add_page_break()
    
    # ============ 9. БЕЗОПАСНОСТЬ ============
    doc.add_heading('9. БЕЗОПАСНОСТЬ И ВНЕШНИЕ ССЫЛКИ', level=1)
    
    doc.add_heading('9.1 Preconnect', level=2)
    
    add_table(doc, ['Домен', 'Тип', 'Статус'], [
        ['images.unsplash.com', 'preconnect + dns-prefetch', '✓'],
        ['experience.tripster.ru', 'preconnect + dns-prefetch', '✓'],
        ['i.pravatar.cc', 'preconnect', '✓'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('9.2 Внешние ссылки', level=2)
    
    add_table(doc, ['Компонент', 'rel="noopener noreferrer"', 'target="_blank"'], [
        ['ContactPage (соцсети)', '✓', '✓'],
        ['PublicFooter', '✗ Внутренние ссылки', '—'],
    ])
    
    doc.add_page_break()
    
    # ============ 10. ДОСТУПНОСТЬ ============
    doc.add_heading('10. ДОСТУПНОСТЬ (A11Y)', level=1)
    
    doc.add_heading('10.1 ARIA-атрибуты', level=2)
    
    add_table(doc, ['Компонент', 'Элемент', 'aria-label', 'Статус'], [
        ['PublicHeader', 'Hamburger menu', '✓ "Открыть меню"', '✓'],
        ['PublicHeader', 'Search icon', '✓ "Открыть поиск"', '✓'],
        ['TourCard', 'Favorite button', '✓ Динамический', '✓'],
        ['TourCard', 'Photo navigation', '✓ "Предыдущее/Следующее фото"', '✓'],
        ['ArticlePage', 'Share button', '✓ "Поделиться"', '✓'],
        ['ArticlePage', 'Bookmark button', '✓ "Сохранить"', '✓'],
        ['HomePage', 'Decorative text', '✓ aria-hidden="true"', '✓'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('10.2 Семантический HTML', level=2)
    
    add_table(doc, ['Элемент', 'Использование'], [
        ['<header>', '✓ PublicHeader'],
        ['<nav>', '✓ Навигация'],
        ['<main>', '✗ Не используется'],
        ['<section>', '✓ Все секции'],
        ['<article>', '✓ ArticlePage'],
        ['<footer>', '✓ PublicFooter'],
        ['<aside>', '✗ Не используется'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('10.3 Структура заголовков', level=2)
    
    add_table(doc, ['Страница', 'H1', 'H2', 'H3', 'Статус'], [
        ['HomePage', '1 (sr-only)', '6', '3', '✓ Единственный H1'],
        ['ToursPage', '1 (динамический)', '2', '1', '✓'],
        ['TourDetailPage', '1 (название тура)', '5', '3', '✓'],
        ['AboutPage', '1', '10', '8', '✓'],
        ['ContactPage', '1', '3', '4', '✓'],
        ['FAQPage', '1', '1', '0', '✓'],
        ['JournalPage', '1', '2', 'Много', '✓'],
        ['ArticlePage', '1', 'Динамически', 'Динамически', '✓'],
    ])
    
    doc.add_page_break()
    
    # ============ 11. PWA ============
    doc.add_heading('11. PWA И MANIFEST', level=1)
    
    doc.add_heading('11.1 Manifest.json', level=2)
    
    add_table(doc, ['Поле', 'Значение', 'Статус'], [
        ['name', 'Inturex — Экскурсии по Азии', '✓'],
        ['short_name', 'Inturex', '✓'],
        ['description', '✓ Присутствует', '✓'],
        ['start_url', '/', '✓'],
        ['display', 'standalone', '✓'],
        ['background_color', '#ffffff', '✓'],
        ['theme_color', '#FF385C', '✓'],
        ['orientation', 'portrait-primary', '✓'],
        ['scope', '/', '✓'],
        ['lang', 'ru', '✓'],
        ['icons', '6 размеров (16, 32, 180, 192, 512, SVG)', '✓'],
        ['categories', '["travel", "tourism"]', '✓'],
        ['prefer_related_applications', 'false', '✓'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('11.2 PWA Meta-теги', level=2)
    
    add_table(doc, ['Тег', 'Значение'], [
        ['mobile-web-app-capable', 'yes'],
        ['apple-mobile-web-app-status-bar-style', 'black-translucent'],
    ])
    
    doc.add_page_break()
    
    # ============ 12. FAVICONS ============
    doc.add_heading('12. FAVICONS', level=1)
    
    add_table(doc, ['Файл', 'Размер', 'Формат', 'Статус'], [
        ['favicon.svg', 'any', 'SVG', '✓'],
        ['favicon.ico', '—', 'ICO', '✓'],
        ['favicon-16x16.png', '16x16', 'PNG', '✓'],
        ['favicon-32x32.png', '32x32', 'PNG', '✓'],
        ['apple-touch-icon.png', '180x180', 'PNG', '✓'],
        ['icon-192.png', '192x192', 'PNG', '✓'],
        ['icon-512.png', '512x512', 'PNG', '✓'],
        ['logo.png', '—', 'PNG', '✓'],
        ['og-image.jpg', '—', 'JPG', '✓'],
    ])
    
    doc.add_page_break()
    
    # ============ 13. ТЕХНИЧЕСКИЕ ФАЙЛЫ ============
    doc.add_heading('13. ТЕХНИЧЕСКИЕ ФАЙЛЫ', level=1)
    
    add_table(doc, ['Файл', 'Расположение', 'Статус'], [
        ['robots.txt', '/public/ + API', '✓'],
        ['sitemap.xml', '/public/ + API', '✓'],
        ['manifest.json', '/public/', '✓'],
        ['favicon.svg', '/public/', '✓'],
        ['favicon.ico', '/public/', '✓'],
        ['og-image.jpg', '/public/', '✓'],
    ])
    
    doc.add_page_break()
    
    # ============ 14. CORE WEB VITALS ============
    doc.add_heading('14. CORE WEB VITALS ФАКТОРЫ', level=1)
    
    doc.add_heading('14.1 Preload/Preconnect', level=2)
    
    add_table(doc, ['Ресурс', 'Тип', 'Статус'], [
        ['images.unsplash.com', 'preconnect + dns-prefetch', '✓'],
        ['experience.tripster.ru', 'preconnect + dns-prefetch', '✓'],
        ['i.pravatar.cc', 'preconnect', '✓'],
        ['Шрифты', '✗ Не preload', '—'],
        ['Критический CSS', '✗ Не preload', '—'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('14.2 Lazy Loading', level=2)
    
    add_table(doc, ['Элемент', 'Статус'], [
        ['Изображения ниже fold', '✓ loading="lazy"'],
        ['Hero изображения', '✗ Lazy (должен быть eager)'],
        ['Компоненты', '✗ Нет React.lazy'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('14.3 Код', level=2)
    
    add_table(doc, ['Аспект', 'Статус'], [
        ['Bundle splitting', '✓ Vite автоматически'],
        ['Tree shaking', '✓ Vite автоматически'],
        ['Минификация', '✓ Production build'],
        ['Source maps', '✗ Не в production'],
    ])
    
    doc.add_page_break()
    
    # ============ 15. ПЕРЕЛИНКОВКА ============
    doc.add_heading('15. ВНУТРЕННЯЯ ПЕРЕЛИНКОВКА', level=1)
    
    doc.add_heading('15.1 Навигация', level=2)
    
    add_table(doc, ['Компонент', 'Ссылки'], [
        ['PublicHeader', 'Главная, Экскурсии, Журнал, О нас, Стать гидом, Войти/Регистрация'],
        ['PublicFooter', '4 города, 4 компания, 3 информация'],
        ['Breadcrumbs', '✓ TourDetailPage, ArticlePage'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('15.2 Контекстные ссылки', level=2)
    
    add_table(doc, ['Страница', 'Ссылки на'], [
        ['HomePage', 'Страны, Города, Рубрики, Туры'],
        ['ToursPage', 'Туры, Города (при выборе страны)'],
        ['AboutPage', 'Туры, Стать гидом'],
        ['JournalPage', 'Статьи, Туры'],
        ['ArticlePage', 'Журнал, Туры по стране'],
    ])
    
    doc.add_page_break()
    
    # ============ 16. ИНДЕКСАЦИЯ ============
    doc.add_heading('16. ИНДЕКСАЦИЯ', level=1)
    
    doc.add_heading('16.1 Robots директивы', level=2)
    
    add_table(doc, ['URL', 'Индексация'], [
        ['/', '✓ Разрешена'],
        ['/tours', '✓ Разрешена'],
        ['/tours/{id}', '✓ Разрешена'],
        ['/journal', '✓ Разрешена'],
        ['/journal/{slug}', '✓ Разрешена'],
        ['/about', '✓ Разрешена'],
        ['/contact', '✓ Разрешена'],
        ['/faq', '✓ Разрешена'],
        ['/dashboard/*', '✗ Закрыта'],
        ['/admin/*', '✗ Закрыта'],
        ['/login', '✗ Закрыта'],
        ['/register', '✗ Закрыта'],
        ['/api/*', '✗ Закрыта'],
    ])
    
    doc.add_paragraph()
    doc.add_heading('16.2 Canonical URLs', level=2)
    
    add_table(doc, ['Страница', 'Canonical', 'Статус'], [
        ['HomePage', '✓ Абсолютный', '✓'],
        ['ToursPage', '✓ Динамический с параметрами', '✓'],
        ['TourDetailPage', '✓ /tours/{id}', '✓'],
        ['AboutPage', '✓ /about', '✓'],
        ['ContactPage', '✓ /contact', '✓'],
        ['FAQPage', '✓ /faq', '✓'],
        ['JournalPage', '✗ Отсутствует', '✗'],
        ['ArticlePage', '✓ /journal/{slug}', '✓'],
    ])
    
    doc.add_page_break()
    
    # ============ 17. СТРУКТУРА URL ============
    doc.add_heading('17. СТРУКТУРА URL', level=1)
    
    add_table(doc, ['Паттерн', 'Пример', 'Статус'], [
        ['Главная', '/', '✓'],
        ['Каталог', '/tours', '✓'],
        ['Фильтр страны', '/tours?location=Таиланд', '✓ (кириллица)'],
        ['Тур', '/tours/{id}', '✓ (числовой ID)'],
        ['Журнал', '/journal', '✓'],
        ['Статья', '/journal/{slug}', '✓ (slug)'],
        ['О нас', '/about', '✓'],
        ['Контакты', '/contact', '✓'],
        ['FAQ', '/faq', '✓'],
        ['Заявка', '/request', '✓'],
        ['Стать гидом', '/become-guide', '✓'],
        ['Условия', '/terms', '✓'],
        ['Политика', '/privacy', '✓'],
    ])
    
    doc.add_page_break()
    
    # ============ 18. КОЛИЧЕСТВЕННЫЕ ПОКАЗАТЕЛИ ============
    doc.add_heading('18. КОЛИЧЕСТВЕННЫЕ ПОКАЗАТЕЛИ', level=1)
    
    add_table(doc, ['Метрика', 'Значение'], [
        ['Страниц в sitemap (статический)', '53'],
        ['Стран', '11'],
        ['Городов в sitemap', '32+'],
        ['Типов JSON-LD схем', '8'],
        ['Favicon размеров', '7'],
        ['Preconnect доменов', '3'],
        ['Вопросов в FAQPage JSON-LD', '18'],
        ['Ссылок в footer', '11'],
        ['Ссылок в header nav', '5'],
    ])
    
    doc.add_page_break()
    
    # ============ 19. СВОДНАЯ ТАБЛИЦА ============
    doc.add_heading('19. СВОДНАЯ ТАБЛИЦА СТАТУСОВ', level=1)
    
    add_table(doc, ['Категория', 'Статус', 'Детали'], [
        ['Meta-теги', '✓ 95%', 'JournalPage без Helmet'],
        ['Open Graph', '✓ 100%', 'Полностью настроены'],
        ['Twitter Cards', '✓ 95%', 'Неправильный атрибут (property vs name)'],
        ['JSON-LD', '✓ 85%', 'JournalPage, ArticlePage без схем'],
        ['Hreflang', '✓ 100%', 'Только RU'],
        ['Robots.txt', '✓ 100%', 'Статический + динамический'],
        ['Sitemap.xml', '✓ 100%', 'Статический + динамический'],
        ['Изображения', '~ 70%', 'Нет WebP, нет srcset'],
        ['Доступность', '✓ 85%', 'Нет <main>, нет skip links'],
        ['PWA', '✓ 100%', 'Полный manifest'],
        ['Favicons', '✓ 100%', 'Все размеры'],
        ['Preload/Preconnect', '~ 60%', 'Нет preload шрифтов'],
        ['Перелинковка', '✓ 90%', 'Хорошая структура'],
        ['Canonical', '✓ 90%', 'JournalPage без canonical'],
    ])
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Итог
    conclusion = doc.add_paragraph()
    conclusion.add_run('ОБЩАЯ ОЦЕНКА: ').bold = True
    conclusion.add_run('SEO настроен на 90%+ для основных страниц. Сайт готов к индексации поисковыми системами.')
    
    # Сохранение
    filename = f'SEO_Audit_Inturex_{datetime.now().strftime("%Y%m%d")}.docx'
    doc.save(filename)
    print(f'Файл сохранён: {filename}')
    return filename

if __name__ == '__main__':
    create_audit()
